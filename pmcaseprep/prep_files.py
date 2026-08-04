"""Turn an uploaded CV or JD into the plain text the Prep Engine already eats.

The paste box was always the real interface — everything downstream
(`extract_units`, `extract_target`) takes a string. This module exists so a
candidate can hand over the PDF they actually have instead of retyping it, and
so the *only* thing that changes downstream is where the string came from.

Deliberately deterministic: no model call, no network. Extraction that costs a
model call would make "attach my CV" a paid, rate-limited, failable step
before the user has seen anything work, and it would double the cost of every
build. pypdf and python-docx are both pure-Python, so the offline test suite
can exercise the whole path.

The one case we cannot serve is a scanned PDF — a photograph of a CV carries
no text layer, and OCR is a different project. We detect that and say so,
rather than handing the engine an empty string and letting it fail later as a
mysterious "no units found".
"""

from __future__ import annotations

import io
import re
from typing import Callable

# A CV is a couple of pages; a JD is less. This is a generous ceiling that
# still refuses the 40MB design portfolio someone will inevitably drag in.
MAX_UPLOAD_BYTES = 5 * 1024 * 1024

# Below this many characters we assume the file had no real text layer rather
# than that someone has a two-line CV. Tuned low on purpose: a false "scanned"
# verdict is worse than letting a genuinely tiny file through to the engine.
_MIN_MEANINGFUL_CHARS = 40

SUPPORTED_SUFFIXES = (".pdf", ".docx", ".txt", ".md", ".markdown", ".text", ".rtf")


class UnsupportedFile(ValueError):
    """The file can't become text. The message is shown to the user verbatim,
    so it must always name the fix, not the failure."""


def _suffix(filename: str) -> str:
    name = (filename or "").strip().lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def _tidy(text: str) -> str:
    """Normalize whitespace without destroying the line structure the
    extractor prompts rely on to see one bullet per line."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")  # PDFs love a non-breaking space
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)  # keep paragraph breaks, drop the rest
    return text.strip()


def _from_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # latin-1 decodes any byte string, so reaching here means the file was
    # not text at all — almost always a renamed binary.
    raise UnsupportedFile("that file isn't readable as text — try a PDF, DOCX, or paste it")


def _from_rtf(data: bytes) -> str:
    """RTF is close enough to text that stripping the control words beats
    refusing the file — Pages and Word both export it by default."""
    raw = data.decode("latin-1", errors="ignore")
    raw = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: chr(int(m.group(1), 16)), raw)
    raw = re.sub(r"\\par[d]?\b", "\n", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)  # remaining control words
    raw = raw.replace("{", "").replace("}", "")
    return raw


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise UnsupportedFile(
            "PDF reading isn't available on this server — paste the text instead"
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty password unlocks the common "print-protected" case.
            try:
                reader.decrypt("")
            except Exception:  # noqa: BLE001
                pass
        pages = [(page.extract_text() or "") for page in reader.pages]
    except UnsupportedFile:
        raise
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedFile(
            "couldn't read that PDF — it may be password-protected. "
            "Try exporting it again, or paste the text instead"
        ) from exc
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise UnsupportedFile(
            "DOCX reading isn't available on this server — paste the text instead"
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedFile(
            "couldn't read that Word file — re-save it as .docx or PDF, "
            "or paste the text instead"
        ) from exc
    parts = [p.text for p in document.paragraphs]
    # CVs are very often laid out as an invisible table; skipping tables would
    # silently return a name and nothing else.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" — ".join(cells))
    return "\n".join(parts)


_READERS: dict[str, Callable[[bytes], str]] = {
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".rtf": _from_rtf,
    ".txt": _from_plain,
    ".text": _from_plain,
    ".md": _from_plain,
    ".markdown": _from_plain,
}


def extract_text(filename: str, data: bytes) -> str:
    """filename + bytes -> plain text, or UnsupportedFile with a fix in it."""
    if not data:
        raise UnsupportedFile("that file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        mb = MAX_UPLOAD_BYTES // (1024 * 1024)
        raise UnsupportedFile(f"that file is over {mb}MB — attach a smaller one")

    suffix = _suffix(filename)
    if suffix == ".doc":
        raise UnsupportedFile(
            "legacy .doc isn't supported — re-save it as .docx or PDF"
        )
    if suffix == ".pages":
        raise UnsupportedFile("export the Pages file to PDF or DOCX first")
    reader = _READERS.get(suffix)
    if reader is None:
        raise UnsupportedFile("attach a PDF, DOCX, or plain text file")

    text = _tidy(reader(data))
    if len(text) < _MIN_MEANINGFUL_CHARS:
        if suffix == ".pdf":
            raise UnsupportedFile(
                "that PDF has no selectable text — it's probably a scan. "
                "Paste the text instead"
            )
        raise UnsupportedFile("couldn't find any text in that file")
    return text
